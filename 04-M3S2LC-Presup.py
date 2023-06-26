import sys
from PyQt5.QtWidgets import QApplication, QWidget, QLabel, QTextEdit, QLineEdit, QPushButton, QVBoxLayout, QMessageBox
from PyQt5.QtGui import QIcon
from docx import Document

class PresupuestoWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Lions Cars - Productos y servicios - Presupuesto")
        self.setWindowIcon(QIcon("logo.jpeg"))
        
        self.title_label = QLabel("Lions Cars")
        self.subtitle_label = QLabel("Prestación de servicio")
        
        self.detalle_label = QLabel("Detalle del servicio:")
        self.detalle_text = QTextEdit()
        
        self.valor_label = QLabel("Valor del servicio:")
        self.valor_edit = QLineEdit()
        
        self.apellido_label = QLabel("Apellido del destinatario:")
        self.apellido_edit = QLineEdit()
        
        self.crear_button = QPushButton("Crear Presupuesto")
        self.crear_button.clicked.connect(self.crear_presupuesto)
        
        layout = QVBoxLayout()
        layout.addWidget(self.title_label)
        layout.addWidget(self.subtitle_label)
        layout.addWidget(self.detalle_label)
        layout.addWidget(self.detalle_text)
        layout.addWidget(self.valor_label)
        layout.addWidget(self.valor_edit)
        layout.addWidget(self.apellido_label)
        layout.addWidget(self.apellido_edit)
        layout.addWidget(self.crear_button)
        self.setLayout(layout)
    
    def crear_presupuesto(self):
        detalle = self.detalle_text.toPlainText()
        valor = self.valor_edit.text()
        apellido = self.apellido_edit.text()
        
        if not detalle or not valor or not apellido:
            QMessageBox.warning(self, "Crear Presupuesto", "Por favor, completa todos los campos.")
            return
        
        document = Document()
        document.add_heading("Presupuesto", level=1)
        document.add_paragraph(f"Apellido del destinatario: {apellido}")
        document.add_paragraph(f"Detalle del servicio: {detalle}")
        document.add_paragraph(f"Valor del servicio: {valor}")
        
        try:
            document.save("Presupuesto.docx")
            QMessageBox.information(self, "Crear Presupuesto", "Se ha creado el presupuesto correctamente.")
        except:
            QMessageBox.warning(self, "Crear Presupuesto", "Ha ocurrido un error al crear el presupuesto.")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    presupuesto_window = PresupuestoWindow()
    presupuesto_window.show()
    sys.exit(app.exec_())
